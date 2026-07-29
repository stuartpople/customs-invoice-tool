"""
Multi-format document extraction module
Supports PDF, Excel, and Word documents
"""
from pdf_extractor import extract_text_from_pdf, parse_line_items, extract_invoice_metadata
from countries import normalize_country_iso, normalize_item_country_fields
import os
import re
import pandas as pd
from typing import List, Dict, Callable, Optional, Tuple
import io

_EXCEL_EXTENSIONS = frozenset({'xlsx', 'xls', 'xlsm'})

# Excel/Word imports
try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False


def _get_api_secret(key_name: str) -> Optional[str]:
    """Streamlit secrets (cloud) then environment variable (local)."""
    try:
        import streamlit as st
        key = st.secrets.get(key_name, "")
        if key:
            return key
    except Exception:
        pass
    return os.environ.get(key_name) or None


def _excel_workbook_to_text(file_obj) -> str:
    """Flatten all sheets to tab-separated text for AI / regex fallback parsing."""
    file_obj.seek(0)
    excel_file = pd.ExcelFile(file_obj)
    parts: list[str] = []
    for sheet_name in excel_file.sheet_names:
        file_obj.seek(0)
        df = pd.read_excel(file_obj, sheet_name=sheet_name, header=None)
        df = df.dropna(how='all')
        if df.empty:
            continue
        parts.append(f"--- SHEET {sheet_name} ---")
        parts.append(df.to_csv(sep='\t', index=False, header=False))
    return '\n'.join(parts)


def _normalize_hs_code(raw, trade_direction: str) -> Optional[str]:
    if raw is None or (isinstance(raw, float) and pd.isna(raw)):
        return None
    if isinstance(raw, (int, float)):
        if isinstance(raw, float) and raw == int(raw):
            code_str = str(int(raw))
        else:
            code_str = str(raw).split('.')[0]
    else:
        code_str = re.sub(r'\D', '', str(raw).strip())
    if not code_str.isdigit() or len(code_str) < 6:
        return None
    if trade_direction.lower() == 'export':
        return code_str[:8]
    if len(code_str) == 8:
        return code_str + '00'
    return code_str


def _normalize_llm_items(llm_items: list, trade_direction: str) -> List[Dict]:
    """Map LLM extractor output to spreadsheet item shape."""
    out: List[Dict] = []
    for it in llm_items or []:
        if not isinstance(it, dict):
            continue
        code = _normalize_hs_code(it.get('commodity_code'), trade_direction)
        value = it.get('value')
        if value is None:
            value = it.get('total_value')
        try:
            value = float(value) if value is not None else None
        except (TypeError, ValueError):
            value = None
        qty = it.get('quantity')
        try:
            qty = float(qty) if qty is not None else None
        except (TypeError, ValueError):
            qty = None
        desc = str(it.get('description') or '').strip()
        if not desc and not code and not value:
            continue
        out.append({
            'commodity_code': code or '',
            'description': desc or (f'Item {code}' if code else 'Line item'),
            'quantity': qty,
            'uom': str(it.get('unit') or 'PCS').strip() or 'PCS',
            'total_value': value,
            'country_of_origin': normalize_country_iso(it.get('country_origin')),
            'net_weight': it.get('net_weight'),
            'currency': str(it.get('currency') or 'GBP').upper()[:3],
            'needs_review': not bool(code),
        })
    return out


def _filter_items_to_source_text(items: List[Dict], source_text: str) -> List[Dict]:
    """Drop rows whose description does not appear in the workbook dump."""
    from line_item_parser import LineItemParser
    return LineItemParser._filter_llm_items_to_source_text(items, source_text)


def _excel_regex_fallback(text: str, trade_direction: str) -> List[Dict]:
    """Regex parsers for sheet text — no AI."""
    from line_item_parser import LineItemParser
    parser = LineItemParser()
    if parser._should_use_bas_parser(text):
        items = parser._parse_bas_commercial_format(text.split("\n"), trade_direction, {})
        items = parser._finalize_parsed_items(items, text)
        if items:
            return _normalize_llm_items(items, trade_direction)
    try:
        items = parse_line_items(text, trade_direction=trade_direction)
        if items:
            normalized = _normalize_llm_items(items, trade_direction)
            return _filter_items_to_source_text(normalized, text)
    except Exception:
        pass
    return []


def _excel_llm_fallback(file_obj, trade_direction: str) -> List[Dict]:
    """When column-based Excel parsing finds nothing, try regex then grounded AI."""
    text = _excel_workbook_to_text(file_obj)
    if len(text.strip()) < 40:
        return []

    regex_items = _excel_regex_fallback(text, trade_direction)
    if regex_items:
        return regex_items

    # BAS-style spreadsheets must not use AI — it invents lines not on the sheet.
    from line_item_parser import LineItemParser
    if LineItemParser._should_use_bas_parser(text):
        return []

    for key_name, extract_fn in (
        ('GOOGLE_API_KEY', 'extract_with_gemini'),
        ('OPENAI_API_KEY', 'extract_with_llm'),
    ):
        api_key = _get_api_secret(key_name)
        if not api_key:
            continue
        try:
            from llm_extractor import extract_with_gemini, extract_with_llm
            fn = extract_with_gemini if extract_fn == 'extract_with_gemini' else extract_with_llm
            llm_items, _meta = fn(text, api_key)
            normalized = _normalize_llm_items(llm_items, trade_direction)
            grounded = _filter_items_to_source_text(normalized, text)
            if grounded:
                return grounded
        except Exception:
            continue
    return []


def _excel_data_items(items: List[Dict]) -> List[Dict]:
    """Line items only — exclude metadata markers and error dicts."""
    return [
        i for i in items
        if isinstance(i, dict) and not i.get('_excel_metadata') and 'error' not in i
    ]


def extract_from_file_with_progress(file_obj, filename: str, trade_direction: str = "export", progress_callback: Optional[Callable] = None):
    """
    Extract data from file with progress updates for OCR.
    
    Args:
        file_obj: File object
        filename: Name of the file
        trade_direction: "export" or "import"
        progress_callback: Function(page_num, total_pages) for progress updates
        
    Returns:
        Tuple of (text, items, metadata)
    """
    file_ext = filename.lower().split('.')[-1]
    
    if file_ext == 'pdf':
        # PDF extraction with progress
        text = extract_text_from_pdf(file_obj, use_ocr=True, progress_callback=progress_callback)
        
        # Parse items and metadata
        items = parse_line_items(text, trade_direction=trade_direction)
        metadata = extract_invoice_metadata(text)
        
        # Add CDS defaults to metadata based on trade direction
        if trade_direction.lower() == 'import':
            metadata['cpc_code'] = '4000'  # Free circulation (default for imports)
            metadata['valuation_method'] = '1'  # Transaction value (Method 1)
        else:  # export
            metadata['cpc_code'] = '1000'  # Permanent export (default for exports)
        
        return text, items, metadata
    else:
        # Other formats don't need progress callbacks
        return extract_from_file(file_obj, filename, trade_direction)


def extract_from_excel(file_obj, trade_direction: str = "export") -> List[Dict]:
    """
    Extract invoice data from Excel file.
    Handles multiple formats including JD Sports, RS Components, etc.
    Also extracts metadata like entry type (B1/H1) and type (E/I).
    
    Args:
        file_obj: File-like object from Streamlit file_uploader
        trade_direction: "export" or "import"
        
    Returns:
        List of line items. First item may be metadata dict if found.
    """
    try:
        # Reset file pointer
        file_obj.seek(0)
        
        # Read Excel file - try all sheets
        excel_file = pd.ExcelFile(file_obj)
        all_items = []
        
        # Extract entry type (B1/H1) and type (E/I) from first few rows
        # These are typically in the header area before data
        excel_metadata = {}
        try:
            first_sheet = excel_file.sheet_names[0] if excel_file.sheet_names else None
            if first_sheet:
                file_obj.seek(0)
                df_header = pd.read_excel(file_obj, sheet_name=first_sheet, nrows=10, header=None)
                
                # Flatten the first 10 rows and search for entry type / type codes
                header_text = ' '.join(str(v) for v in df_header.values.flatten() if pd.notna(v))
                
                # Search for entry type (B1 or H1)
                if ' B1 ' in f' {header_text} ':
                    excel_metadata['entry_type'] = 'B1'
                    excel_metadata['direction_detected'] = 'export'
                elif ' H1 ' in f' {header_text} ':
                    excel_metadata['entry_type'] = 'H1'
                    excel_metadata['direction_detected'] = 'import'
                else:
                    # Try finding them without spaces
                    if 'B1' in header_text:
                        excel_metadata['entry_type'] = 'B1'
                    if 'H1' in header_text:
                        excel_metadata['entry_type'] = 'H1'
                
                # Search for type (E for export, I for import)
                if ' E ' in f' {header_text} ':
                    excel_metadata['type'] = 'E'
                elif ' I ' in f' {header_text} ':
                    excel_metadata['type'] = 'I'
                
                # If metadata found, store as first item with special marker
                if excel_metadata:
                    all_items.append({
                        '_excel_metadata': True,
                        **excel_metadata
                    })
        except Exception:
            pass  # Continue without metadata if extraction fails
        
        for sheet_name in excel_file.sheet_names:
            file_obj.seek(0)
            df = pd.read_excel(file_obj, sheet_name=sheet_name, header=None)
            
            # Skip sheets with very few rows (likely summary/cover sheets)
            if len(df) < 2:
                continue
            
            # --- Auto-detect header row ---
            # Scan first 40 rows for one that looks like a column header
            # (contains keywords like HS CODE, DESCRIPTION, QTY, etc.)
            header_keywords = {'hs code', 'commodity code', 'commodity', 'tariff',
                               'description', 'goods description', 'goods desc',
                               'qty', 'quantity', 'value', 'value (gbp)', 'weight', 'origin',
                               'uom', 'unit cost', 'line total', 'total', 'hs',
                               'part', 'sku', 'article', 'item', 'amount', 'price',
                               'net', 'gross', 'coo', 'line no', 'line #'}
            header_row = None
            max_scan = min(40, len(df))
            
            for i in range(max_scan):
                row_vals = [str(v).lower().strip() for v in df.iloc[i] if pd.notna(v)]
                matches = sum(1 for v in row_vals if any(kw in v for kw in header_keywords))
                if matches >= 2:  # At least 2 recognised column keywords
                    header_row = i
                    break
            
            if header_row is not None:
                # Re-read with detected header row
                file_obj.seek(0)
                df = pd.read_excel(file_obj, sheet_name=sheet_name, header=header_row)
            else:
                # Fall back to first row as header (original behaviour)
                file_obj.seek(0)
                df = pd.read_excel(file_obj, sheet_name=sheet_name)
            
            # Drop completely empty rows
            df = df.dropna(how='all')
            
            # Normalise column names for matching
            col_lower = {col: str(col).lower().strip() for col in df.columns}
            
            # Invoice sheets usually have value + qty; allow value-only layouts (some templates).
            _value_check = _find_columns(df.columns, col_lower, [
                ['line total', 'line_total', 'line value'],
                ['total value', 'total_value', 'invoice value'],
                ['extended', 'amount'],
                ['value', 'price'],
                ['total', 'cost', 'amount'],
            ])
            _qty_check = _find_columns(df.columns, col_lower, [
                ['qty', 'quantity'],
                ['units', 'unit'],
            ])
            # Fallback: if value column not found by name, check unnamed columns for numeric data
            # (handles formats like Rhenus where value/price column has no header text)
            if not _value_check and _qty_check:
                for _col in df.columns:
                    _cl = col_lower.get(_col, str(_col).lower())
                    if not _cl.startswith('unnamed:'):
                        continue
                    _nums = pd.to_numeric(df[_col], errors='coerce').dropna()
                    if len(_nums) >= max(3, len(df) * 0.4) and _nums.mean() > 0.01:
                        _value_check = [_col]
                        break
            if not _value_check:
                continue
            
            # --- Column detection with priority ordering ---
            # HS / Commodity code columns
            code_columns = _find_columns(df.columns, col_lower, [
                ['hs code', 'hs_code'],                              # exact first
                ['commodity code', 'commodity_code', 'comm code'],
                ['tariff'],
                ['hs', 'commodity', 'classification'],               # partial
            ])
            
            # Description columns
            desc_columns = _find_columns(df.columns, col_lower, [
                ['description', 'goods desc', 'goods description'],  # exact first
                ['desc', 'product', 'item name', 'item description'],
                ['item', 'name'],                                    # partial fallback
            ])
            
            # Quantity columns
            qty_columns = _find_columns(df.columns, col_lower, [
                ['qty', 'quantity'],
                ['count'],
            ])
            
            # Value columns — prefer "line total" over generic "total" or "cost"
            value_columns = _find_columns(df.columns, col_lower, [
                ['line total', 'line_total', 'line value'],          # exact line total
                ['total value', 'total_value', 'invoice value'],
                ['value', 'price'],
                ['total', 'cost', 'amount'],                         # less specific
            ])
            # Fallback: unnamed column with numeric data (e.g. Rhenus format)
            if not value_columns:
                _skip = set(c for cols in [code_columns, desc_columns, qty_columns] for c in cols)
                for _col in df.columns:
                    if _col in _skip:
                        continue
                    _cl = col_lower.get(_col, str(_col).lower())
                    if not _cl.startswith('unnamed:'):
                        continue
                    _nums = pd.to_numeric(df[_col], errors='coerce').dropna()
                    if len(_nums) >= max(3, len(df) * 0.4) and _nums.mean() > 0.01:
                        value_columns = [_col]
                        break
            
            # Weight columns — prefer line/total weight over unit weight.
            # 'Line Weight' is the per-row total (unit_weight * qty) and is
            # correct even when unit_weight rounds to 0 for very light items.
            weight_columns = _find_columns(df.columns, col_lower, [
                ['total (kg)', 'total_kg', 'total weight', 'net weight', 'nett weight', 'line weight'],
                ['net_weight', 'weight (kg)', 'weight_kg', 'line_weight'],
                ['weight', 'net', 'gross', 'kg'],
            ])
            
            country_columns = _find_columns(df.columns, col_lower, [
                ['country of destination', 'country of origin', 'destination', 'origin country'],
                ['origin', 'country', 'coo', 'c.o.o', 'coc'],
            ])
            # Fallback: unnamed column where values look like 2-3 letter country codes
            if not country_columns:
                _used = set(value_columns)
                for _col in df.columns:
                    if _col in _used:
                        continue
                    _cl = col_lower.get(_col, str(_col).lower())
                    if not _cl.startswith('unnamed:'):
                        continue
                    _vals = df[_col].dropna().astype(str).str.strip()
                    _cc_like = _vals.str.match(r'^[A-Z]{2,3}$').sum()
                    if _cc_like >= max(3, len(_vals) * 0.5):
                        country_columns = [_col]
                        break
            
            # UOM columns
            uom_columns = _find_columns(df.columns, col_lower, [
                ['uom', 'unit of measure'],
                ['unit'],
            ])
            
            # Material columns (for JD Sports - may contain useful info)
            material_columns = _find_columns(df.columns, col_lower, [
                ['material'],
            ])
            
            # Extract rows — stop at end-of-data boundary
            consecutive_empty = 0
            max_empty_gap = 3  # Stop after 3 consecutive rows with no HS code
            
            for idx, row in df.iterrows():
                # --- End-of-data detection ---
                # Check if this is a TOTAL / summary row (signals end of items)
                row_text = ' '.join(str(v).lower() for v in row if pd.notna(v))
                if any(marker in row_text for marker in ['total value', 'grand total', 'invoice total', 'sub total', 'subtotal']):
                    break
                
                # Try to find commodity code
                commodity_code = None
                if code_columns:
                    for col in code_columns:
                        val = row[col]
                        if pd.notna(val):
                            commodity_code = _normalize_hs_code(val, trade_direction)
                            if commodity_code:
                                break
                
                # Extract other fields (needed before row skip / description-only path)
                description = None
                if desc_columns:
                    for col in desc_columns:
                        if pd.notna(row[col]):
                            description = str(row[col]).strip()
                            break
                
                quantity = None
                if qty_columns:
                    for col in qty_columns:
                        if pd.notna(row[col]):
                            try:
                                quantity = float(_clean_numeric(row[col]))
                                break
                            except:
                                pass
                
                value = None
                if value_columns:
                    for col in value_columns:
                        if pd.notna(row[col]):
                            try:
                                value = float(_clean_numeric(row[col]))
                                break
                            except:
                                pass
                
                weight = None
                if weight_columns:
                    for col in weight_columns:
                        if pd.notna(row[col]):
                            try:
                                weight = float(_clean_numeric(row[col]))
                                break
                            except:
                                pass
                
                country = None
                if country_columns:
                    for col in country_columns:
                        if pd.notna(row[col]):
                            country = normalize_country_iso(str(row[col]).strip())
                            break
                
                uom = None
                if uom_columns:
                    for col in uom_columns:
                        if pd.notna(row[col]):
                            uom = str(row[col]).strip()
                            break
                
                # Skip phantom rows: must have a value > 0
                if not value or value <= 0:
                    if not commodity_code:
                        consecutive_empty += 1
                        if consecutive_empty >= max_empty_gap:
                            break
                    continue

                if not commodity_code:
                    # Commercial invoice rows: description + value but no HS column yet
                    if not (description and len(description) > 2):
                        consecutive_empty += 1
                        if consecutive_empty >= max_empty_gap:
                            break
                        continue
                    consecutive_empty = 0
                    all_items.append({
                        "commodity_code": "",
                        "description": description,
                        "quantity": quantity,
                        "uom": uom or "PCS",
                        "total_value": value,
                        "country_of_origin": country or "",
                        "net_weight": weight,
                        "currency": "GBP",
                        "needs_review": True,
                    })
                    continue

                # Valid coded row — reset gap counter
                consecutive_empty = 0

                all_items.append({
                    "commodity_code": commodity_code,
                    "description": description if description else f"Item {commodity_code}",
                    "quantity": quantity,
                    "uom": uom or "PCS",
                    "total_value": value,
                    "country_of_origin": country or "",
                    "net_weight": weight,
                    "currency": "GBP",
                    "needs_review": False,
                })

        for item in all_items:
            if isinstance(item, dict) and not item.get('_excel_metadata'):
                normalize_item_country_fields(item)

        data_items = _excel_data_items(all_items)
        if not data_items:
            file_obj.seek(0)
            fallback_items = _excel_llm_fallback(file_obj, trade_direction)
            if fallback_items:
                if all_items and all_items[0].get('_excel_metadata'):
                    return [all_items[0]] + fallback_items
                return fallback_items

        return all_items

    except Exception as e:
        file_obj.seek(0)
        try:
            fallback_items = _excel_llm_fallback(file_obj, trade_direction)
            if fallback_items:
                return fallback_items
        except Exception:
            pass
        return [{"error": f"Error reading Excel: {str(e)}"}]


def _clean_numeric(val) -> str:
    """Strip currency symbols, commas, and whitespace from a value for float conversion."""
    s = str(val).strip()
    # Remove common currency symbols and thousands separators
    for ch in ['£', '$', '€', '¥', ',', '\u00a3', '\u20ac']:
        s = s.replace(ch, '')
    return s.strip()


def _find_columns(columns, col_lower_map, priority_groups):
    """
    Find matching columns using prioritised keyword groups.
    Returns the first group that matches any columns.
    Each group is a list of keywords — columns are matched if they
    contain any keyword (case-insensitive).
    """
    for keywords in priority_groups:
        matches = []
        for col in columns:
            cl = col_lower_map.get(col, str(col).lower())
            for kw in keywords:
                if kw in cl:
                    matches.append(col)
                    break
        if matches:
            return matches
    return []


def _word_document_to_text(doc) -> str:
    """Flatten Word paragraphs and table cells for metadata/debugging."""
    parts = [para.text for para in doc.paragraphs if para.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            parts.append(' | '.join(cell.text.replace('\n', ' / ') for cell in row.cells))
    return '\n'.join(parts)


def _word_decimal(raw) -> Optional[float]:
    """Parse GBP values using either decimal point or decimal comma."""
    match = re.search(r'-?\d[\d\s,.]*', str(raw or '').replace('\xa0', ' '))
    if not match:
        return None
    value = re.sub(r'\s+', '', match.group(0))
    if ',' in value and '.' not in value:
        value = value.replace(',', '.')
    elif ',' in value and '.' in value:
        # Last separator is decimal; the other is thousands.
        if value.rfind(',') > value.rfind('.'):
            value = value.replace('.', '').replace(',', '.')
        else:
            value = value.replace(',', '')
    try:
        return float(value)
    except ValueError:
        return None


def _parse_word_commercial_tables(doc, trade_direction: str) -> List[Dict]:
    """Extract standard commercial-invoice rows directly from Word tables."""
    items: List[Dict] = []
    suspicious_decimal_comma: List[Dict] = []

    for table in doc.tables:
        if not table.rows:
            continue
        headers = [
            re.sub(r'\s+', ' ', cell.text).strip().lower()
            for cell in table.rows[0].cells
        ]

        def column(*names):
            for idx, header in enumerate(headers):
                if any(name in header for name in names):
                    return idx
            return None

        qty_col = column('quantity', 'qty')
        desc_col = column('description', 'goods')
        country_col = column('country of origin', 'origin')
        hs_col = column('hs code', 'commodity code', 'tariff')
        value_col = column('value', 'amount', 'line total')
        if None in (qty_col, desc_col, hs_col, value_col):
            continue

        for row in table.rows[1:]:
            cells = [re.sub(r'\s+', ' ', cell.text).strip() for cell in row.cells]
            if max(qty_col, desc_col, hs_col, value_col) >= len(cells):
                continue

            description = cells[desc_col]
            code = _normalize_hs_code(cells[hs_col], trade_direction)
            qty = _word_decimal(cells[qty_col])
            value = _word_decimal(cells[value_col])
            if not description or not code or qty is None or value is None:
                # Repeated headers and package separator rows land here.
                continue
            if qty <= 0 or value <= 0:
                continue

            country = ''
            if country_col is not None and country_col < len(cells):
                country = normalize_country_iso(cells[country_col])
            item = {
                'commodity_code': code,
                'description': description,
                'quantity': qty,
                'uom': 'PCS',
                'total_value': value,
                'country_of_origin': country,
                'net_weight': None,
                'currency': 'GBP',
                'needs_review': False,
            }
            normalize_item_country_fields(item)
            items.append(item)
            raw_value = cells[value_col]
            if ',' in raw_value and '.' not in raw_value:
                suspicious_decimal_comma.append(item)

    # Surface a real inconsistency instead of silently treating comma as thousands.
    paragraph_text = '\n'.join(para.text for para in doc.paragraphs)
    printed_match = re.search(
        r'\bTotal\s+value\s*:\s*[£$€]?\s*([\d,.]+)',
        paragraph_text,
        re.IGNORECASE,
    )
    printed_total = _word_decimal(printed_match.group(1)) if printed_match else None
    row_total = round(sum(float(item['total_value']) for item in items), 2)
    if printed_total is not None and abs(row_total - printed_total) >= 0.01:
        difference = round(row_total - printed_total, 2)
        candidates = [
            item for item in suspicious_decimal_comma
            if abs(float(item['total_value']) - abs(difference)) < 0.01
        ]
        review_items = candidates or suspicious_decimal_comma
        note = (
            f"Invoice rows total GBP {row_total:.2f}, but the printed total is "
            f"GBP {printed_total:.2f} (difference GBP {difference:.2f}). "
            "Check the comma-decimal value on this row."
        )
        for item in review_items:
            item['needs_review'] = True
            item['review_notes'] = note

    return items


def extract_from_word(file_obj, trade_direction: str = "export") -> List[Dict]:
    """
    Extract invoice data from Word document.
    
    Args:
        file_obj: File-like object from Streamlit file_uploader
        trade_direction: "export" or "import"
        
    Returns:
        List of line items
    """
    if not DOCX_AVAILABLE:
        return [{"error": "python-docx library not installed"}]
    
    try:
        # Reset file pointer
        file_obj.seek(0)
        
        # Read Word document
        doc = Document(file_obj)
        
        # Word tables preserve columns, so use them before flattening text.
        items = _parse_word_commercial_tables(doc, trade_direction)
        if items:
            return items

        text = _word_document_to_text(doc)
        
        # Use the PDF parser logic (it works on text)
        from pdf_extractor import parse_line_items
        items = parse_line_items(text, trade_direction=trade_direction)
        
        return items
        
    except Exception as e:
        return [{"error": f"Error reading Word document: {str(e)}"}]


def extract_from_file(file_obj, filename: str, trade_direction: str = "export") -> Tuple[str, List[Dict], Dict]:
    """
    Extract data from uploaded file based on file type.
    
    Args:
        file_obj: File-like object from Streamlit file_uploader
        filename: Name of the file
        trade_direction: "export" or "import"
        
    Returns:
        Tuple of (extracted_text, list_of_items, metadata_dict)
    """
    try:
        # Reset file pointer at start
        file_obj.seek(0)
        
        file_ext = filename.lower().split('.')[-1]
        
        # Default metadata based on trade direction
        if trade_direction.lower() == 'import':
            metadata = {
                'cpc_code': '4000',  # Free circulation
                'valuation_method': '1',  # Transaction value
                'incoterm': None,
                'currency': 'GBP',
                'total_invoice_value': None,
                'total_gross_weight': None,
                'total_net_weight': None,
                'number_of_packages': None,
                'package_type': None,
                'invoice_number': None,
                'invoice_date': None,
            }
        else:  # export
            metadata = {
                'cpc_code': '1000',  # Permanent export
                'incoterm': None,
                'currency': 'GBP',
                'total_invoice_value': None,
                'total_gross_weight': None,
                'total_net_weight': None,
                'number_of_packages': None,
                'package_type': None,
                'invoice_number': None,
                'invoice_date': None,
            }
        
        if file_ext == 'pdf':
            text = extract_text_from_pdf(file_obj)
            items = parse_line_items(text, trade_direction=trade_direction)
            metadata.update(extract_invoice_metadata(text))
            return text, items, metadata
        
        elif file_ext in _EXCEL_EXTENSIONS:
            items = extract_from_excel(file_obj, trade_direction)
            # Create a text representation for debug view
            text = f"Excel file: {filename}\nExtracted {len(items)} items from spreadsheet"
            return text, items, metadata
        
        elif file_ext in ['docx', 'doc']:
            items = extract_from_word(file_obj, trade_direction)
            file_obj.seek(0)
            doc = Document(file_obj)
            text = _word_document_to_text(doc)
            metadata.update(extract_invoice_metadata(text))

            invoice_ref = re.search(r'\bINV\s*:\s*([A-Z0-9/-]+)', text, re.IGNORECASE)
            invoice_date = re.search(
                r'\b(\d{1,2}/\d{1,2}/\d{4})\b', text, re.IGNORECASE
            )
            incoterm = re.search(r'\bINCO\s*Terms?\s+([A-Z]{3})\b', text, re.IGNORECASE)
            printed_total = re.search(
                r'\bTotal\s+value\s*:\s*[£$€]?\s*([\d,.]+)', text, re.IGNORECASE
            )
            gross_total = re.search(
                r'\bTotal\s+Weight\s*:\s*([\d,.]+)\s*kg', text, re.IGNORECASE
            )
            package_numbers = [
                int(n) for n in re.findall(r'\bPackage\s+(\d+)\s*:', text, re.IGNORECASE)
            ]
            net_weights = [
                _word_decimal(n) for n in re.findall(
                    r'\bNett\s+Weight\s*:\s*([\d,.]+)\s*kg', text, re.IGNORECASE
                )
            ]
            if invoice_ref:
                metadata['invoice_number'] = invoice_ref.group(1)
            if invoice_date:
                metadata['invoice_date'] = invoice_date.group(1)
            if incoterm:
                metadata['incoterm'] = incoterm.group(1).upper()
            if printed_total:
                metadata['total_invoice_value'] = _word_decimal(printed_total.group(1))
            if gross_total:
                metadata['total_gross_weight'] = _word_decimal(gross_total.group(1))
            if package_numbers:
                metadata['number_of_packages'] = max(package_numbers)
                metadata['package_type'] = 'PK'
            valid_net_weights = [n for n in net_weights if n is not None]
            if valid_net_weights:
                metadata['total_net_weight'] = round(sum(valid_net_weights), 3)
            return text, items, metadata
        
        else:
            return f"Unsupported file type: {file_ext}", [], metadata
            
    except Exception as e:
        error_msg = f"Error processing {filename}: {str(e)}"
        return error_msg, [], metadata
